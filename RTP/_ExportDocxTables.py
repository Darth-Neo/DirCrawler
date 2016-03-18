#
# Crawl a directory for documents and pull out the text
#
import os
import sys
from docx import Document

import unicodedata
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from nl_lib.Logger import *
logger = setupLogging(__name__)
logger.setLevel(DEBUG)


def docxText(df, of):

    with open(of, "w") as f:
        document = Document(df)

        lp = list()

        n = 0

        parent_elm = document.element.body

        for child in parent_elm.iterchildren():

            if isinstance(child, CT_P):
                nl = list()
                nl.append(n)
                nl.append(u"P")
                nl.append(child)
                lp.append(nl)
                logger.info(u"Paragraph %s" % child.text)

            elif isinstance(child, CT_Tbl):
                nl = list()
                nl.append(n)
                nl.append(u"T")
                nl.append(child)
                lp.append(nl)
                logger.info(u"Table     %s" % child.text)

            n += 1

        if False:
            paragraphs = document.paragraphs
            for x in paragraphs:
                logger.info(u"%s" % x.text)

            sections = document.sections
            for x in sections:
                logger.info(u"%s" % x)

        tables = document.tables
        nt = 0
        for table in tables:

            nt += 1
            nc = table._column_count

            nr = 0
            for row in table.rows:
                nr += 1

                rs = u""
                try:
                    cl = list()

                    for n in range(0, nc):
                        try:

                            txt = u"%s" % row.cells[n].paragraphs[0].text
                            rs += txt + u","

                            logger.debug(u"%d => %s" % (n, txt))

                        except Exception, msg:
                            logger.info(u"Opps %s" % msg)

                    rs = unicodedata.normalize(u'NFKD', rs).encode(u'ascii', u'ignore')
                    f.write(rs + os.linesep)

                except Exception, msg:
                    logger.error(u"%s" % msg)

                logger.debug(u"    row = %d" % nr)

            logger.info(u"table = %d" % nt)

            f.write(os.linesep)
            f.write("New Table%s" % os.linesep)
            f.write(os.linesep)


if __name__ == u'__main__':
    # Set the directory you want to start from
    df = os.getcwd() + os.sep + u"batchspecificationv14.docx"
    of = "tables_14.csv"
    docxText(df, of)
